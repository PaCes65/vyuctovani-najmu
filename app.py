import streamlit as st
import pandas as pd
from pypdf import PdfReader
import google.generativeai as genai
import json
import re
import io
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.error("Kritická závislost: Chybí modul python-docx. Přidejte jej do requirements.txt.")
    st.stop()

# ==========================================
# 1. KONFIGURACE A INIT STAVU
# ==========================================
st.set_page_config(page_title="Vyúčtování PRO", layout="wide", page_icon="🏢")

if "GEMINI_API_KEY" not in st.secrets:
    st.error("Chyba konfigurace: GEMINI_API_KEY nenalezen v Secrets.")
    st.stop()

genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

DEFAULT_DATA = {
    "profil": {"poskytovatel": "", "platce": "", "adresa": "", "byt": "", "typ": ""},
    "osa_najmu": [{"od_mesice": 1, "do_mesice": 12, "najem": 0, "zaloha": 0}],
    "osoby": [{"jmeno": "", "od_mesice": 1, "do_mesice": 12}],
    "lonsko": {"vysledek": 0, "typ": "Zadne"},
    "naklady": {"svj": [], "dalsi_sluzby": []},
    "vypocty": {
        "suma_transakci": 0.0, "celkova_pohledavka_najem": 0.0, "korekce_saldo": 0.0,
        "disponibilni_zalohy": 0.0, "uznatelne_naklady": 0.0, "saldo_konecne": 0.0,
        "pocet_transakci": 0, "filtr_transakci": "", "prumerna_obsazenost": 0.0
    }
}

if "db" not in st.session_state:
    st.session_state.db = json.loads(json.dumps(DEFAULT_DATA))

# Pomocná funkce pro smazání stavu widgetů
def clear_ui_states(keys):
    for k in keys:
        if k in st.session_state:
            del st.session_state[k]

# ==========================================
# 2. IO & AI SUBSYSTÉM
# ==========================================
@st.cache_resource
def zjisti_model():
    try:
        modely = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        for pref in ['models/gemini-1.5-flash', 'models/gemini-2.0-flash-exp', 'models/gemini-pro']:
            if pref in modely: return pref
        return modely[0] if modely else None
    except: return None

def ai_volani(prompt):
    model_name = zjisti_model()
    try:
        model = genai.GenerativeModel(model_name, generation_config={"response_mime_type": "application/json"})
        res = model.generate_content(prompt)
        return json.loads(res.text)
    except Exception as e:
        try:
            model_fb = genai.GenerativeModel(model_name)
            res_fb = model_fb.generate_content(prompt)
            match = re.search(r'\{.*\}', res_fb.text, re.DOTALL)
            if match: return json.loads(match.group())
            raise Exception("Selhání regex fallbacku.")
        except Exception as e_fb:
            raise Exception(f"AI Chyba: {e_fb}")

def cteni_pdf(file_obj):
    try:
        reader = PdfReader(file_obj)
        text = "".join(p.extract_text() + "\n" for p in reader.pages if p.extract_text())
        return text
    except Exception as e:
        raise Exception(f"Chyba dekódování PDF: {e}")

def generovat_word_protokol(db):
    doc = Document()
    title = doc.add_heading('PROTOKOL O ZÚČTOVÁNÍ ZÁLOH NA SLUŽBY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header = doc.add_paragraph()
    p_header.add_run(f"Dle zákona č. 67/2013 Sb., o službách\n").italic = True
    p_header.add_run(f"Generováno: {datetime.now().strftime('%d.%m.%Y')}")
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('I. Identifikace subjektů a předmětu', level=1)
    t_id = doc.add_table(rows=4, cols=2)
    t_id.style = 'Table Grid'
    t_id.columns[0].width = Inches(2.0)
    r = t_id.rows
    r[0].cells[0].text = 'Předmět:'; r[0].cells[1].text = f"{db['profil']['adresa']}, Byt: {db['profil']['byt']}"
    r[1].cells[0].text = 'Klasifikace:'; r[1].cells[1].text = db['profil']['typ']
    r[2].cells[0].text = 'Poskytovatel:'; r[2].cells[1].text = db['profil']['poskytovatel']
    r[3].cells[0].text = 'Plátce:'; r[3].cells[1].text = db['profil']['platce']
    
    v = db['vypocty']
    doc.add_heading('II. Evidenční počet osob (obsazenost)', level=1)
    doc.add_paragraph("Přehled uživatelů bytu rozhodný pro rozúčtování nákladů podle zákona č. 67/2013 Sb.:")
    t_osoby = doc.add_table(rows=1, cols=3)
    t_osoby.style = 'Light Shading'
    hr_o = t_osoby.rows[0].cells
    hr_o[0].text = 'Jméno'; hr_o[1].text = 'Období užívání'; hr_o[2].text = 'Osoboměsíců'
    for os in db.get("osoby", []):
        row = t_osoby.add_row().cells
        pmesicu = max(0, (os.get("do_mesice", 12) - os.get("od_mesice", 1)) + 1)
        row[0].text = os.get("jmeno", "Neuvedeno")
        row[1].text = f"Měsíc {os.get('od_mesice', 1)} až {os.get('do_mesice', 12)}"
        row[2].text = str(pmesicu)
    doc.add_paragraph(f"\nPrůměrný roční počet osob v bytě: {v.get('prumerna_obsazenost', 0.0):.2f}").bold = True

    doc.add_heading('III. Časová osa předpisu plnění', level=1)
    t_predpis = doc.add_table(rows=1, cols=4)
    t_predpis.style = 'Light Shading'
    hr = t_predpis.rows[0].cells
    hr[0].text = 'Měsíce'; hr[1].text = 'Délka (m)'; hr[2].text = 'Nájem (CZK)'; hr[3].text = 'Sumace Nájem (CZK)'
    for o in db["osa_najmu"]:
        row = t_predpis.add_row().cells
        pocet_m = max(0, (o.get("do_mesice", 12) - o.get("od_mesice", 1)) + 1)
        sub_najem = pocet_m * o.get("najem", 0)
        row[0].text = f"{o.get('od_mesice', 1)} až {o.get('do_mesice', 12)}"; row[1].text = str(pocet_m)
        row[2].text = f"{o.get('najem', 0):,.2f}"; row[3].text = f"{sub_najem:,.2f}"
    doc.add_paragraph(f"\nSrážka předpisu úhrad celkem: {v['celkova_pohledavka_najem']:,.2f} CZK").bold = True

    doc.add_heading('IV. Finanční syntéza a disponibilní zálohy', level=1)
    doc.add_paragraph(f"1. Připsané transakce (N={v['pocet_transakci']}): {v['suma_transakci']:,.2f} CZK")
    doc.add_paragraph(f"2. Aplikace srážky předpisu (Pol. III): - {v['celkova_pohledavka_najem']:,.2f} CZK")
    doc.add_paragraph(f"3. Aplikace salda N-1 ({db['lonsko']['typ']}): {'+' if v['korekce_saldo'] >= 0 else ''}{v['korekce_saldo']:,.2f} CZK")
    doc.add_paragraph(f"Disponibilní zálohy k vyrovnání = {v['disponibilni_zalohy']:,.2f} CZK").bold = True

    doc.add_heading('V. Zúčtovatelné náklady objektu', level=1)
    agregovane = db['naklady']['svj'] + db['naklady']['dalsi_sluzby']
    uznatelne = [p for p in agregovane if p.get('preuctovatelne')]
    if uznatelne:
        t_naklady = doc.add_table(rows=1, cols=2)
        t_naklady.style = 'Table Grid'
        t_naklady.rows[0].cells[0].text = 'Identifikátor nákladu'; t_naklady.rows[0].cells[1].text = 'Hodnota (CZK)'
        for n in uznatelne:
            row = t_naklady.add_row().cells
            row[0].text = n['nazev']; row[1].text = f"{n['castka']:,.2f}"
        doc.add_paragraph(f"\nUznatelné náklady k tíži plátce celkem: {v['uznatelne_naklady']:,.2f} CZK").bold = True
    else:
        doc.add_paragraph("Kalkulace neobsahuje nákladové položky.")

    doc.add_heading('VI. Konečné vypořádání salda', level=1)
    saldo = v['saldo_konecne']
    p_res = doc.add_paragraph()
    p_res.add_run(f"VÝSLEDNÉ KONEČNÉ SALDO: {saldo:,.2f} CZK\n").bold = True
    
    if saldo > 0:
        p_res.add_run(f"\nKlasifikace: PŘEPLATEK ({abs(saldo):,.2f} CZK).\n").bold = True
        p_res.add_run("Identifikován závazek poskytovatele vůči plátci s povinností vypořádání.")
    elif saldo < 0:
        p_res.add_run(f"\nKlasifikace: NEDOPLATEK ({abs(saldo):,.2f} CZK).\n").bold = True
        p_res.add_run("Identifikována splatná pohledávka poskytovatele za plátcem dle § 7 zák. č. 67/2013 Sb.")
    else:
        p_res.add_run("\nKlasifikace: VYROVNÁNO. Finanční závazky nulové.").bold = True

    doc.add_paragraph("\n\nSmluvní strany stvrzují podpisem akceptaci kalkulační metodiky a výsledného salda.")
    doc.add_paragraph('\n\nV ........................................ dne ........................................')
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.columns[0].width = Inches(3.0); sig_table.columns[1].width = Inches(3.0)
    sig_table.rows[0].cells[0].text = '..........................................................'
    sig_table.rows[0].cells[1].text = '..........................................................'
    sig_table.rows[1].cells[0].text = f"Za Poskytovatele:\n{db['profil']['poskytovatel']}"
    sig_table.rows[1].cells[1].text = f"Za Plátce:\n{db['profil']['platce']}"
    for r in sig_table.rows:
        for c in r.cells: c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# 3. UŽIVATELSKÉ ROZHRANÍ
# ==========================================
st.title("🏢 Vyúčtování PRO: Komplexní systém s agregací")

t1, t2, t3, t4, t5 = st.tabs(["1. Profil & Osy", "2. Saldo N-1", "3. Náklady N", "4. Agregace & Export", "💾 Perzistence"])

with t1:
    st.subheader("Extrakce struktur a časových řad")
    c_p1, c_p2 = st.columns(2)
    sml_1_pdf = c_p1.file_uploader("Smlouva 1 (Výchozí platnost)", type="pdf", key="file_sml_1")
    sml_2_pdf = c_p2.file_uploader("Smlouva 2 / Dodatek (Změna plnění)", type="pdf", key="file_sml_2")
    
    if st.button("Spustit sémantickou analýzu", type="primary", key="btn_nlp_smlouvy"):
        if not sml_1_pdf: st.warning("Absentuje primární dokument.")
        else:
            with st.spinner("Probíhá sémantická extrakce..."):
                try:
                    txt = f"--- DOKUMENT 1 (Výchozí stav) ---\n{cteni_pdf(sml_1_pdf)}\n"
                    if sml_2_pdf: txt += f"--- DOKUMENT 2 (Změna) ---\n{cteni_pdf(sml_2_pdf)}\n"
                    
                    prompt = f"""
                    Jsi právní auditor. Extrahuj z textů JSON pro 12-ti měsíční rok (1-12).
                    
                    PRAVIDLA EXTRAKCE:
                    1. JMENNÉ ROZLIŠENÍ:
                       - Vyhledej přesná jména a příjmení osob.
                       - Podnájemní smlouva: Poskytovatel = Nájemce, Plátce = Podnájemce.
                       - Nájemní smlouva: Poskytovatel = Pronajímatel, Plátce = Nájemce.
                    2. OSOBY: Seznam jmen všech osob v bytě (hlavní plátce + spolubydlící/manželka).
                    3. ČASOVÁ MATICE (osa_najmu) - STRIKTNĚ:
                       - Smlouva 1 platí od 1. měsíce.
                       - Smlouva 2 (dodatek) platí od M. měsíce.
                       - Pokud je v textu Smlouva 2, MUSÍŠ rozdělit rok do DVOU záznamů v poli osa_najmu.
                         Záznam A: od 1 do (M-1) s původními cenami.
                         Záznam B: od M do 12 s novými cenami.
                       - 'najem' (čistý nájem), 'zaloha' (na služby).
                    
                    VZOR SCHÉMATU VÝSTUPU:
                    {{
                        "profil": {{"typ": "Nájemní smlouva", "poskytovatel": "Karel Novák", "platce": "Jan Dvořák", "adresa": "Adresa", "byt": "Číslo"}},
                        "osa_najmu": [
                            {{"od_mesice": 1, "do_mesice": 5, "najem": 10000, "zaloha": 3000}},
                            {{"od_mesice": 6, "do_mesice": 12, "najem": 11000, "zaloha": 3500}}
                        ],
                        "osoby": [
                            {{"jmeno": "Jan Dvořák", "od_mesice": 1, "do_mesice": 12}},
                            {{"jmeno": "Jana Dvořáková", "od_mesice": 1, "do_mesice": 12}}
                        ]
                    }}
                    TEXT K ANALÝZE: {txt}
                    """
                    res = ai_volani(prompt)
                    if res:
                        st.session_state["posledni_json_smlouvy"] = res
                        if "profil" in res and isinstance(res["profil"], dict):
                            st.session_state.db["profil"].update(res["profil"])
                        if "osa_najmu" in res and res["osa_najmu"]:
                            st.session_state.db["osa_najmu"] = res["osa_najmu"]
                        if "osoby" in res and res["osoby"]:
                            st.session_state.db["osoby"] = res["osoby"]
                        
                        # Vynucený flush widgetů
                        clear_ui_states(["inp_posk", "inp_plat", "inp_adr", "inp_byt", "editor_osa_najmu", "editor_osoby"])
                        st.rerun()
                except Exception as e: st.error(f"Chyba NLP procesingu: {e}")

    # Vizualizace surového JSON
    if "posledni_json_smlouvy" in st.session_state:
        with st.expander("🔍 Inspekce surových dat vyčtených AI (Smlouvy)", expanded=False):
            st.json(st.session_state["posledni_json_smlouvy"])

    c_i = st.columns(4)
    st.session_state.db["profil"]["poskytovatel"] = c_i[0].text_input("Poskytovatel", st.session_state.db["profil"].get("poskytovatel", ""), key="inp_posk")
    st.session_state.db["profil"]["platce"] = c_i[1].text_input("Plátce", st.session_state.db["profil"].get("platce", ""), key="inp_plat")
    st.session_state.db["profil"]["adresa"] = c_i[2].text_input("Adresa", st.session_state.db["profil"].get("adresa", ""), key="inp_adr")
    st.session_state.db["profil"]["byt"] = c_i[3].text_input("ID Bytu", st.session_state.db["profil"].get("byt", ""), key="inp_byt")
    
    col_osa, col_osoby = st.columns(2)
    
    with col_osa:
        st.markdown("#### Matice časové osy předpisů")
        st.session_state.db["osa_najmu"] = st.data_editor(st.session_state.db["osa_najmu"], num_rows="dynamic", key="editor_osa_najmu")
        
        st.markdown("##### 🧮 Agregace předpisu (Dílčí a celkové součty)")
        agg_osa = []
        celk_najem, celk_zaloha = 0.0, 0.0
        for o in st.session_state.db["osa_najmu"]:
            m_count = max(0, (o.get("do_mesice", 12) - o.get("od_mesice", 1)) + 1)
            n_sum = m_count * float(o.get("najem", 0))
            z_sum = m_count * float(o.get("zaloha", 0))
            celk_najem += n_sum
            celk_zaloha += z_sum
            agg_osa.append({
                "Interval": f"{o.get('od_mesice')}. - {o.get('do_mesice')}. měsíc",
                "Měsíců": m_count,
                "Suma Nájemné": f"{n_sum:,.2f} CZK",
                "Suma Zálohy": f"{z_sum:,.2f} CZK"
            })
        st.dataframe(pd.DataFrame(agg_osa), use_container_width=True, hide_index=True)
        st.info(f"**CELKOVÝ PŘEDPIS ROKU:** Nájemné: **{celk_najem:,.2f} CZK** | Zálohy: **{celk_zaloha:,.2f} CZK**")
    
    with col_osoby:
        st.markdown("#### Evidenční list osob")
        st.session_state.db["osoby"] = st.data_editor(st.session_state.db["osoby"], num_rows="dynamic", key="editor_osoby")
        
        st.markdown("##### 🧮 Agregace obsazenosti")
        total_osobomesicu = sum(max(0, (o.get("do_mesice", 12) - o.get("od_mesice", 1)) + 1) for o in st.session_state.db.get("osoby", []))
        prumerna_obsazenost = total_osobomesicu / 12.0
        st.session_state.db["vypocty"]["prumerna_obsazenost"] = prumerna_obsazenost
        st.info(f"**Kumulativní osoboměsíce:** {total_osobomesicu}\n\n**Průměrná roční obsazenost bytu:** {prumerna_obsazenost:.2f} osob")

with t2:
    st.subheader("Historické saldo (N-1)")
    l_pdf = st.file_uploader("Vyúčtování (N-1)", type="pdf", key="file_lon_pdf")
    if l_pdf and st.button("Extrakce salda", key="btn_saldo_lonsko"):
        res = ai_volani(f"Najdi saldo. Přeplatek (kladná), Nedoplatek (záporná). Schéma JSON: {{'vysledek': int, 'typ': 'Preplatek/Nedoplatek'}}. Text: {cteni_pdf(l_pdf)}")
        if res: 
            st.session_state["posledni_json_lonsko"] = res
            st.session_state.db["lonsko"].update(res)
            clear_ui_states(["sel_typ_lonsko", "num_vysledek_lonsko"])
            st.rerun()
            
    if "posledni_json_lonsko" in st.session_state:
        with st.expander("🔍 Inspekce extrakce (Saldo N-1)", expanded=False): st.json(st.session_state["posledni_json_lonsko"])
    
    c_l = st.columns(2)
    st.session_state.db["lonsko"]["typ"] = c_l[0].selectbox("Klasifikátor salda", ["Preplatek", "Nedoplatek", "Zadne"], index=["Preplatek", "Nedoplatek", "Zadne"].index(st.session_state.db["lonsko"].get("typ", "Zadne")), key="sel_typ_lonsko")
    st.session_state.db["lonsko"]["vysledek"] = c_l[1].number_input("Absolutní hodnota (CZK)", value=abs(float(st.session_state.db["lonsko"].get("vysledek", 0))), key="num_vysledek_lonsko")

with t3:
    st.subheader("Rozvrh nákladových středisek (N)")
    svj_pdf = st.file_uploader("Primární náklad (SVJ)", type="pdf", key="file_svj_pdf")
    if svj_pdf and st.button("Dekompozice položek", key="btn_dek_svj"):
        res = ai_volani(f"Rozřaď položky. 'preuctovatelne'=true pro služby (voda, teplo, výtah). Schéma JSON: {{'polozky': [{{'nazev':'string', 'castka':100, 'preuctovatelne':true}}]}}. Text: {cteni_pdf(svj_pdf)}")
        if res and "polozky" in res: 
            st.session_state["posledni_json_svj"] = res
            st.session_state.db["naklady"]["svj"] = res["polozky"]
            clear_ui_states(["editor_naklady_svj"])
            st.rerun()
            
    if "posledni_json_svj" in st.session_state:
        with st.expander("🔍 Inspekce extrakce (SVJ)", expanded=False): st.json(st.session_state["posledni_json_svj"])
    
    st.session_state.db["naklady"]["svj"] = st.data_editor(st.session_state.db["naklady"]["svj"], num_rows="dynamic", key="editor_naklady_svj")
    
    sum_svj_vse = sum(float(p.get("castka", 0)) for p in st.session_state.db["naklady"].get("svj", []))
    sum_svj_uzn = sum(float(p.get("castka", 0)) for p in st.session_state.db["naklady"].get("svj", []) if p.get("preuctovatelne"))
    st.info(f"**SVJ Celkem (Hrubý úhrn):** {sum_svj_vse:,.2f} CZK | **Neuznatelné (Fond oprav aj.):** {(sum_svj_vse - sum_svj_uzn):,.2f} CZK | **Přeúčtovatelné na plátce:** {sum_svj_uzn:,.2f} CZK")
    
    st.divider()
    c_ds = st.columns([3, 1])
    d_pdf = c_ds[0].file_uploader("Sekundární náklad (Plyn/El.)", type="pdf", key="file_sek_pdf")
    typ_d = c_ds[1].selectbox("Identifikátor", ["Elektřina", "Plyn", "Internet", "Ostatní"], key="sel_typ_sek")
    if d_pdf and st.button("Přidat entitu", key="btn_add_sek"):
        res = ai_volani(f"Agreguj fakturovanou sumu. Schéma JSON: {{'castka': int}}. Text: {cteni_pdf(d_pdf)}")
        if res: 
            st.session_state["posledni_json_sek"] = res
            st.session_state.db["naklady"]["dalsi_sluzby"].append({"nazev": typ_d, "castka": float(res.get("castka", 0)), "preuctovatelne": True})
            clear_ui_states(["editor_dalsi_sluzby"])
            st.rerun()
            
    if "posledni_json_sek" in st.session_state:
        with st.expander("🔍 Inspekce extrakce (Sekundární náklady)", expanded=False): st.json(st.session_state["posledni_json_sek"])
        
    st.session_state.db["naklady"]["dalsi_sluzby"] = st.data_editor(st.session_state.db["naklady"]["dalsi_sluzby"], num_rows="dynamic", key="editor_dalsi_sluzby")
    
    sum_sek_vse = sum(float(p.get("castka", 0)) for p in st.session_state.db["naklady"].get("dalsi_sluzby", []))
    sum_sek_uzn = sum(float(p.get("castka", 0)) for p in st.session_state.db["naklady"].get("dalsi_sluzby", []) if p.get("preuctovatelne"))
    st.info(f"**Sekundární náklady celkem:** {sum_sek_vse:,.2f} CZK | **Přeúčtovatelné:** {sum_sek_uzn:,.2f} CZK")
    
    st.markdown(f"### **Celkový úhrn uznatelných nákladů (Debet plátce): {(sum_svj_uzn + sum_sek_uzn):,.2f} CZK**")

with t4:
    st.subheader("Finanční agregátor & Výstup")
    filtr = st.text_input("Transakční filtr (Substring matche)", st.session_state.db["profil"].get("platce", ""), key="inp_filtr_trans")
    b_file = st.file_uploader("Export transakcí (XLSX/CSV)", type=["xlsx", "csv"], key="file_banka_xls")
    
    if st.button("KALKULACE & ZÁPIS", type="primary", use_container_width=True, key="btn_calc"):
        if b_file:
            try:
                if b_file.name.endswith('.csv'):
                    df = pd.read_csv(b_file, sep=None, engine='python', header=None, dtype=str)
                else:
                    df = pd.read_excel(b_file, header=None, dtype=str)
                
                mask = df.apply(lambda r: r.astype(str).str.contains(filtr, case=False, na=False).any(), axis=1)
                filtered = df[mask]
                
                if filtered.empty:
                    raise Exception(f"Filtrační argument '{filtr}' nevrátil žádné odpovídající záznamy.")
                
                def parse_fin(val):
                    if pd.isna(val): return 0.0
                    v = str(val).strip()
                    if re.match(r'^\d{1,4}[-\.]\d{1,2}[-\.]\d{1,4}', v): return 0.0
                    v = v.replace(' ', '').replace('\xa0', '')
                    if v.endswith(',-') or v.endswith('.-'): v = v[:-2]
                    v = re.sub(r'[^\d,\.-]', '', v)
                    try:
                        idx_c = v.rfind(',')
                        idx_d = v.rfind('.')
                        if idx_c > idx_d: v = v.replace('.', '').replace(',', '.')
                        else: v = v.replace(',', '')
                        return float(v)
                    except: return 0.0

                sums = {c: filtered[c].apply(parse_fin).sum() for c in filtered.columns}
                target = max(sums, key=sums.get)
                
                if sums[target] == 0:
                    raise Exception("Matice neobsahuje validní numerická finanční data.")
                
                v = st.session_state.db["vypocty"]
                v["filtr_transakci"] = filtr
                v["suma_transakci"] = sums[target]
                v["pocet_transakci"] = len(filtered)
                
                v["celkova_pohledavka_najem"] = sum(max(0, (o.get("do_mesice", 12) - o.get("od_mesice", 1)) + 1) * float(o.get("najem", 0)) for o in st.session_state.db["osa_najmu"])
                
                ls = st.session_state.db["lonsko"]
                v["korekce_saldo"] = float(ls.get("vysledek", 0)) if ls.get("typ") == "Preplatek" else -abs(float(ls.get("vysledek", 0))) if ls.get("typ") == "Nedoplatek" else 0
                
                v["disponibilni_zalohy"] = v["suma_transakci"] - v["celkova_pohledavka_najem"] + v["korekce_saldo"]
                v["uznatelne_naklady"] = sum(float(p.get('castka', 0)) for p in st.session_state.db["naklady"].get("svj", []) + st.session_state.db["naklady"].get("dalsi_sluzby", []) if p.get('preuctovatelne'))
                v["saldo_konecne"] = v["disponibilni_zalohy"] - v["uznatelne_naklady"]
                
                st.session_state["posledni_df_transakce"] = filtered
                st.session_state["posledni_target_col"] = target
                st.success("Kalkulace validována a zapsána do stavové databáze.")
            except Exception as e:
                st.error(f"Kritická chyba procesingu: {e}")

    if st.session_state.db["vypocty"].get("pocet_transakci", 0) > 0:
        if "posledni_df_transakce" in st.session_state:
            st.markdown("#### Matice identifikovaných transakcí")
            st.dataframe(st.session_state["posledni_df_transakce"], use_container_width=True)
            st.info(f"**Identifikován sloupec s hodnotami (Index {st.session_state['posledni_target_col']}). Agregovaná suma transakcí:** {st.session_state.db['vypocty']['suma_transakci']:,.2f} CZK (Počet záznamů: {st.session_state.db['vypocty']['pocet_transakci']})")

        v_ref = st.session_state.db['vypocty']
        st.markdown("---")
        st.markdown(f"**Disponibilní zálohy:** {v_ref.get('disponibilni_zalohy', 0):,.2f} CZK | **Uznatelné náklady:** {v_ref.get('uznatelne_naklady', 0):,.2f} CZK")
        st.metric("Vypočtené konečné saldo", f"{v_ref.get('saldo_konecne', 0):,.2f} CZK")
        
        docx = generovat_word_protokol(st.session_state.db)
        fn = v_ref.get("filtr_transakci", "export").replace(" ", "_")
        st.download_button("📄 EXPORTOVAT DOCX PROTOKOL", data=docx, file_name=f"Zuctovani_{fn}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="btn_export_docx")

with t5:
    st.subheader("Serializace databáze")
    st.download_button("📥 Uložit JSON Image", data=json.dumps(st.session_state.db, indent=4, ensure_ascii=False), file_name="Data_Export.json", mime="application/json", key="btn_export_json")
    st.divider()
    up = st.file_uploader("Obnovit z JSON", type="json", key="file_import_json")
    if up and st.button("RESTORE", key="btn_restore_json"): 
        st.session_state.db = json.load(up)
        st.rerun()
